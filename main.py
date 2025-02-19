from kivy.lang import Builder
from kivy.properties import StringProperty, NumericProperty
from kivy.uix.screenmanager import Screen
from kivymd.app import MDApp
from kivymd.uix.menu import MDDropdownMenu
from kivymd.uix.button import MDIconButton
from kivymd.uix.card import MDCard
from kivymd.uix.label import MDLabel
from kivymd.uix.relativelayout import MDRelativeLayout
from kivymd.uix.gridlayout import MDGridLayout
from kivymd.uix.list import MDList, ThreeLineIconListItem, IconLeftWidget, OneLineIconListItem
from kivy.metrics import dp
from kivymd.uix.fitimage import FitImage
from kivymd.uix.dialog import MDDialog
from kivymd.uix.button import MDRaisedButton
import mysql.connector as connector
import configparser
from docx import Document

conn = connector.connect(
    host="127.0.0.1",
    user="root",
    password="Highpeople11",
    database="gradingrpa", 
    auth_plugin='mysql_native_password'
)

cursor = conn.cursor()
config = configparser.ConfigParser()
config.read('config.ini')
key = config.get('ENCRYPTION', 'key')

class Classes(Screen):
    pass
class Report(Screen):
    pass
class SignIn(Screen):
    pass
class MainScreen(Screen):
    pass

class CourseScreen(Screen):
    pass

class Grades(Screen):
    pass

class Student(Screen):
    pass

class IconListItem(OneLineIconListItem):
    icon = StringProperty

class ReportGen(MDApp):
    def build(self):
        self.theme_cls.primary_palette = 'Green'
        self.sm = Builder.load_file('reportgen.kv')
        self.menu = MDDropdownMenu()
        menu_items = [
            {
                "viewclass": "IconListItem",
                "icon": "git",
                "text": f"Term 1",
                "height": dp(56),
                "on_release": lambda x='term1': self.drop_handling(x)
            },
            {
                "viewclass": "IconListItem",
                "icon": "git",
                "text": f"Term 2",
                "height": dp(56),
                "on_release": lambda x='term2': self.drop_handling(x)
            },
            {
                "viewclass": "IconListItem",
                "icon": "git",
                "text": f"Term 3",
                "height": dp(56),
                "on_release": lambda x='term3': self.drop_handling(x)
            }
            ]
        self.menu = MDDropdownMenu(
            caller=self.sm.get_screen('report').ids.term,
            items=menu_items,
            position="bottom",
            width_mult=4,
        )
        self.menu.bind()
        return self.sm
    
    def drop_handling(self, term):
        print('HANDLE')
        self.menu.dismiss()
        self.dropdown_item.text = term
        

    def sign_in(self, id, password):
        
        if id and password:
            cursor.execute('''SELECT * FROM teacher WHERE id= %s AND AES_DECRYPT(password, %s) = %s ''', (id, key, password, ))
            credentials = cursor.fetchall()
            if credentials:
                self.teacher = credentials[0][0]
                self.teacherfname = credentials[0][1] 
                self.teacherlname = credentials[0][2]
                self.email = credentials[0][3]
                self.number = credentials[0][4]
                self.get_main_screen('courses')
                
            else:
                self.sm.get_screen('signin').ids.teacherid.helper_text =  'ID and/or Password May Be Wrong, Please Try again'
                self.sm.get_screen('signin').ids.password.helper_text =  'ID and/or Password May Be Wrong, Please Try again'
                self.sm.get_screen('signin').ids.teacherid.error = True
                self.sm.get_screen('signin').ids.password.error = True
        else:
            self.sm.get_screen('signin').ids.teacherid.helper_text =  'Both Fields Must Be Completed'
            self.sm.get_screen('signin').ids.password.helper_text =  'Both Fields Must Be Completed'
            self.sm.get_screen('signin').ids.teacherid.error = True
            self.sm.get_screen('signin').ids.password.error = True
    def get_main_screen(self, screen):
        if screen == 'courses':
            # self.sm.get_screen('main').ids.mainwelcome.text = f'Hi {self.teacherfname}, here are the subjects\n you are teaching this term...'
            self.sm.get_screen('main').ids.scroll_main.clear_widgets()
            cursor.execute('''SELECT name, id, class FROM course where teacher=%s''', (self.teacher,))
            coursedetails = cursor.fetchall()
            print('course', coursedetails)
            details = []
            for course in coursedetails:
                ide = course[2]
                cursor.execute('''SELECT name FROM class where id=%s''', (ide, ))
                new = course + cursor.fetchall()[0]
                details.append(new)
            self.sm.get_screen('main').ids.scroll_main.add_widget(self.create_cards(details))
            self.sm.current = 'main'
        elif screen == 'classes':
            self.sm.get_screen('class').ids.scroll_main.clear_widgets()
            cursor.execute('''Select * from class where teacher=%s''', (self.teacher,))
            self.data_table = self.create_list([(item[0], item[1], ' ') for item in cursor.fetchall()])
            self.sm.get_screen('class').ids.scroll_main.add_widget(self.data_table)
            self.sm.current = 'class'
            pass
        self.sm.get_screen('main').ids.nav_drawer.set_state('closed')
        self.sm.get_screen('student').ids.nav_drawer.set_state('closed')

    def create_list(self, row_data):
        """Create an MDList widget."""
        mdlist = MDList()
        for thing in row_data:
            iteration = ThreeLineIconListItem(
                                            IconLeftWidget(icon='signal'),
                                            text= thing[1], 
                                            secondary_text= 'ID: ' + str(thing[0]),
                                            tertiary_text= str(thing[-1]),
                                            divider_color= (0,0,0,1),
                                            bg_color = (0,0,0,0.1),
                                            on_release= lambda x: self.on_row_press(x.text, x.secondary_text, x.tertiary_text),
                                            )
            mdlist.add_widget(iteration)
        return mdlist
    
    def create_cards(self, row_data):
        
        grid = MDGridLayout(cols=4,
                            padding= "50dp",
                            spacing= "50dp",
                            size_hint_y=None,  # Ensures it expands beyond screen
                            size_hint_x = 1,
                            height = 1500,
                            )

        # Center the grid_layout inside the scrollview
        grid.pos_hint = {"center_x": 0.5}
        for course in row_data:
            print(course)
            card = MDCard(
                    MDRelativeLayout(
                        FitImage(
                        source="photos/stock.jpg",
                        size_hint_y=0.6,
                        size_hint_x= 1,
                        pos_hint={"top": 1},
                        radius=(dp(16), dp(16), 0, 0),
                        ),
                        MDIconButton(
                            icon="dots-vertical",
                            pos_hint={"top": 1, "right": 1}
                        ),
                        MDLabel(
                            text=course[0] ,
                            adaptive_size=True,
                            pos=("12dp", "54dp"),
                            bold = True,
                            font_size= '100sp'
                        ),
                        MDLabel(
                            text=course[-1] + f' ({course[2]})',
                            adaptive_size=True,
                            pos=("12dp", "35dp"),
                            id='classname'
                        ),
                         MDLabel(
                            text='Subject ID: ' + str(course[1]),
                            adaptive_size=True,
                            pos=("12dp", "18dp"),
                            id='subjectid'
                        ),
                    ),
                    style='elevated',
                    padding="4dp",
                    size_hint=(None, None),
                    size=("260dp", "300dp"),
                    ripple_behavior=True,
                    pos_hint= {'center_x': 0.5,'center_y': 0.5}, 
                    on_release= lambda instance: self.on_card_press( instance.children[0].ids['subjectid'].text, instance.children[0].ids['classname'].text)
                )
            grid.add_widget(card)
            
        return grid

    def get_all_items(self, table, cond, val):
        cursor.execute(f"SELECT * FROM {table} WHERE {cond}={val}")
        return cursor.fetchall()


    
    def update_grade(self, termval, studentid, courseid):
        grade = self.dialog.content_cls.ids.term.text.strip()   

        if grade:
            print(grade)
            cursor.execute(f'''UPDATE grade SET {termval}=%s WHERE studentid=%s AND courseid=%s''', (grade, studentid, courseid, ))
            conn.commit()
            
        else:
            pass
        inputid = 't' + termval.split('m')[1]
        self.sm.get_screen('student').ids[inputid].text = grade + '%'
        self.dialog.dismiss()
        
    
    def on_card_press(self, id, classname):
        id = id.split(':')[1].strip()
        classid = classname.split('(')[1].split(')')[0]
        print(id, classname)
        self.screen = self.sm.current_screen
        print(self.screen.name)
        if self.screen.name == 'main':
            cursor.execute('''SELECT s.id, fname, lname, term1, term2, term3
                            from gradingrpa.student as s join gradingrpa.course as c
                            on s.class = c.class join gradingrpa.grade as g on c.id = g.courseid
                            WHERE g.courseid = %s and s.class = %s''', (id, classid))
            grade = cursor.fetchall()
            row_input = [(item[0], item[1] + ' ' + item[2], item[3], item[4], item[5], id) for item in grade]
            print(row_input)
            table = self.create_list(row_input)
            self.sm.get_screen(self.sm.current).ids.scroll_main.clear_widgets()
            self.sm.get_screen('course').ids.scroll_course.clear_widgets()
            self.sm.get_screen('course').ids.nav_drawer.set_state('close')
            self.sm.get_screen('course').ids.scroll_course.add_widget(table)
            self.sm.current = 'course'
    
    def on_row_press(self, name, id, courseid):
        self.screen = self.sm.current_screen
        id = id.split(':')[1].strip()
        print(id, courseid)
        self.courseid = courseid
        if self.screen.name == 'course':
            cursor.execute('''SELECT s.id, fname, lname, term1, term2, term3, s.class 
                           FROM gradingrpa.student as s join gradingrpa.grade as g on s.id = g.studentid 
                           WHERE s.id= %s AND g.courseid=%s''', (id, courseid, ))
            details = cursor.fetchall()
            print(details)
            self.studentid = details[0][0]
            self.studentname = details[0][1] + ' ' + details[0][2]
            self.studentclass = details[0][-1]
            t1 = str(details[0][3])
            t2 = str(details[0][4])
            t3 = str(details[0][5])
            self.sm.get_screen('student').ids.studentname.text = self.studentname
            self.sm.get_screen('student').ids.classid.text = 'Class ID: ' + str(self.studentclass)
            self.sm.get_screen('student').ids.t1.text = t1 + '%'
            self.sm.get_screen('student').ids.t2.text = t2 + '%'
            self.sm.get_screen('student').ids.t3.text = t3 + '%'
            self.sm.current = 'student'
        elif self.screen.name == 'class':
            self.content_cls = Report()
            self.dropdown_item = self.content_cls.ids.term
            self.dropdown_item.text = 'Select Term'
            self.dialog = MDDialog(
            title= f'Generate Reports for Class {id}?',
            type='custom',
            content_cls= self.content_cls,
            buttons=[
                MDRaisedButton(
                    text="Yes",
                    on_release=lambda x: self.generate_reports(id)
                ),
                MDRaisedButton(
                    text="No",
                    on_release=lambda x: self.dialog.dismiss()
                ),
                ],
            )
            self.dialog.open()
    def edit_grade(self, term, termval):
        self.dialog = MDDialog(
            title=f'Edit {term} Grades',
            type="custom",
            content_cls=Grades(),
            buttons=[
                MDRaisedButton(
                    text="Submit",
                    on_release=lambda x: self.update_grade(termval, self.studentid, self.courseid)
                ),
                MDRaisedButton(
                    text="CLOSE",
                    on_release=lambda x: self.dialog.dismiss()
                ),
                ],
            )
        self.dialog.open()

    def generate_reports(self, clas):
        cursor.execute('''SELECT * FROM STUDENT WHERE class=%s''', (clas, ))
        students = cursor.fetchall()
        for student in students:
            student_id = student[0]
            student_name = student[1] + ' ' + student[2]
            term = self.dialog.content_cls.ids.term.text.strip()
            # Create a new Document
            doc = Document()

            # Title
            doc.add_heading(f'Report Card for {student_name}', 0)

            # Add student details
            doc.add_paragraph(f'Student Name: {student_name}')
            
            # Add table for grades
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Set column names
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Subject'
            hdr_cells[1].text = 'Grade'

            print(term)
            cursor.execute(f'''SELECT g.{term}, c.name FROM grade as g join course as c on g.courseid = c.id WHERE studentid=%s''', (student_id, ))
            courses = cursor.fetchall()

            for course in courses:
                grade = f'{course[0]}'
                
                subject = course[1]
                

                # Add rows for each subject
 
                row_cells = table.add_row().cells
                row_cells[0].text = subject
                row_cells[1].text = grade

                # Add final remarks
            doc.add_paragraph('\nFinal Remarks:')
            doc.add_paragraph('Excellent progress in most subjects. Keep up the good work!')

            # Save the document
            doc.save(f'reports/{student_name}_{term}_report_card.docx')
        self.dialog.dismiss()


if __name__ == "__main__":
    ReportGen().run()
